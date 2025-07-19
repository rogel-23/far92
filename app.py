import streamlit as st
import pandas as pd
import os
from datetime import datetime
import json
from io import BytesIO
import math
from datetime import datetime, date
from google_drive_utils import upload_to_drive
from streamlit import secrets
import uuid

folder_id = st.secrets.get("1VRY2gQlO2lcI5AiKft2m75-bWUkWqB_Z", None) or "1VRY2gQlO2lcI5AiKft2m75-bWUkWqB_Z"  # fallback

# === Initialisation session_state ===
if "far_arbitres" not in st.session_state:
    st.session_state["far_arbitres"] = []

# === FONCTIONS DE CHARGEMENT/SAUVEGARDE ===
APP_DIR = os.path.dirname(os.path.abspath(__file__))
FILENAME = os.path.join(APP_DIR, "far_arbitres.xlsx")
os.makedirs("rapports", exist_ok=True)

def load_arbitres():
    if os.path.exists(FILENAME):
        df = pd.read_excel(FILENAME)
        for col in ["Rassemblements"]:
            if col not in df.columns:
                df[col] = ""
        return df.to_dict(orient="records")
    else:
        df = pd.DataFrame(columns=[
            "Nom", "Pr√©nom", "Cat√©gorie", "Date de naissance", "√Çge",
            "Club", "T√©l√©phone", "Email", "Rassemblements"
        ])
        df.to_excel(FILENAME, index=False)
        return []


def save_arbitres(data):
    df = pd.DataFrame(data)
    df.to_excel(FILENAME, index=False)

# === TITRE + BOUTON SAUVEGARDE ALIGN√âS ===
col_title, col_save = st.columns([5, 1])

with col_title:
    st.title("‚öΩ FAR 92 - Application de gestion")
    st.markdown("Bienvenue sur l'application officielle de la **Fili√®re Arbitrage R√©gionale du District 92**.")

with col_save:
    st.write("")  # d√©calage vertical
    df = pd.DataFrame(st.session_state["far_arbitres"])

    # Test : ex√©cution locale ou non
    is_local = os.path.exists(".git") or os.getenv("STREAMLIT_ENV") != "cloud"

    if is_local:
        if st.button("üíæ Sauvegarder"):
            df.to_excel("far_arbitres.xlsx", index=False)
            st.success("Fichier Excel sauvegard√© dans le dossier du projet.")
    else:
        buffer = BytesIO()
        df.to_excel(buffer, index=False)
        buffer.seek(0)
        st.download_button(
            "üì• T√©l√©charger Excel",
            data=buffer,
            file_name="far_arbitres.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# === CHARGEMENT INITIAL D'UN FICHIER EXCEL ===
if "fichier_source" not in st.session_state:
    st.session_state["fichier_source"] = None

if st.session_state["fichier_source"] is None:
    st.subheader("üìÇ Charger un fichier Excel FAR")
    uploaded_file = st.file_uploader("S√©lectionnez un fichier Excel", type=["xlsx"])

    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file)
        for col in ["Rassemblements"]:
            if col not in df.columns:
                df[col] = ""
        st.session_state["far_arbitres"] = df.to_dict(orient="records")
        st.session_state["fichier_source"] = uploaded_file.name
        st.success("Fichier charg√© avec succ√®s.")
        st.rerun()
    else:
        st.stop()


# === SESSION STATE ===
if "far_arbitres" not in st.session_state:
    st.session_state["far_arbitres"] = load_arbitres()

# === MENU ===
st.subheader("üìö Modules disponibles")
col1, col2, col3 = st.columns(3)

with col1:
    action = st.radio("Menu", [
    "üìù Compte-rendu rassemblement",
    "üìä R√©capitulatif des rassemblements",
    "üìù Saisie des examens",
    "üìä R√©capitulatif des examens",
    "üõë Ajouter des manquements",
    "üìâ R√©capitulatif des manquements",
    "üìé D√©poser un rapport d'observation",
    "üë§ Fiche arbitre",
    "‚ûï Ajouter / ‚ùå Supprimer un arbitre"
])

# === AJOUT / SUPPRESSION ===
if action == "‚ûï Ajouter / ‚ùå Supprimer un arbitre":
    st.subheader("‚ûï Ajouter un arbitre")
    with st.form("ajout_arbitre_form"):
        nom = st.text_input("Nom").upper()
        prenom = st.text_input("Pr√©nom")
        categorie = st.selectbox("Cat√©gorie FAR", ["FAR-S1", "FAR-S2", "FAR-A1", "FAR-J1", "FAR-F1"])
        date_naissance = st.date_input("Date de naissance", min_value=date(1900, 1, 1))
        club = st.text_input("Club")
        tel = st.text_input("T√©l√©phone")
        email = st.text_input("Adresse mail")
        submit = st.form_submit_button("Ajouter")

        if submit and nom and prenom:
            today = datetime.date.today()
            age = today.year - date_naissance.year - ((today.month, today.day) < (date_naissance.month, date_naissance.day))
            st.session_state["far_arbitres"].append({
                "Nom": nom, "Pr√©nom": prenom, "Cat√©gorie": categorie,
                "Date de naissance": date_naissance.strftime("%d/%m/%Y"), "√Çge": age,
                "Club": club, "T√©l√©phone": tel, "Email": email,
                "Rassemblements": ""
            })
            save_arbitres(st.session_state["far_arbitres"])
            st.success("Ajout√© avec succ√®s.")

    st.subheader("‚ùå Supprimer un arbitre")
    for i, a in enumerate(st.session_state["far_arbitres"]):
        col1, col2, col3, col4 = st.columns([3, 3, 3, 1])
        col1.markdown(f"**{a['Pr√©nom']} {a['Nom']}**")
        col2.write(a["Cat√©gorie"])
        col3.write(a["Club"])
        if col4.button("üóëÔ∏è", key=f"del_{i}"):
            st.session_state["far_arbitres"].pop(i)
            save_arbitres(st.session_state["far_arbitres"])
            st.rerun()

# === EXPORT ===
elif action == "üìÑ Exporter la liste modifi√©e":
    st.subheader("üìÑ Export")
    df = pd.DataFrame(st.session_state["far_arbitres"])
    buffer = BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    st.download_button("üì• T√©l√©charger le fichier", data=buffer, file_name="far_arbitres.xlsx")

# === COMPTE-RENDU RASSEMBLEMENT ===
elif action == "üìù Compte-rendu rassemblement":
    st.subheader("üìù Nouveau compte-rendu")
    type_rass = st.selectbox("Type de rassemblement", ["R√©union", "Stage", "Test physique", "Autre"])

    if type_rass in ["R√©union", "Autre"]:
        with st.form("form_reunion"):
            nom_rass = st.text_input("Nom de la r√©union")
            date_rass = st.date_input("Date de la r√©union")
            statuts = {}
            st.markdown("### Pr√©sence des arbitres")
            commentaires = {}
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Pr√©nom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")
                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Pr√©sent", "Absent excus√©", "Absent non excus√©"], key=f"r_statut_{i}")
                commentaire = col2.text_input("Commentaire individuel", key=f"r_comment_{i}")
                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()

            observations = st.text_area("Observations g√©n√©rales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_rass:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_rass]

                    rass.append({
                        "Nom": nom_rass,
                        "Type": type_rass,
                        "Date": date_rass.strftime("%d/%m/%Y"),
                        "Statut": statuts[f"{a['Pr√©nom']} {a['Nom']}"],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Pr√©nom']} {a['Nom']}", "")
                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("R√©union enregistr√©e avec succ√®s.")

    elif type_rass == "Stage":
        with st.form("form_stage"):
            nom_stage = st.text_input("Nom du stage")
            date_debut = st.date_input("Date de d√©but")
            date_fin = st.date_input("Date de fin")
            statuts = {}
            st.markdown("### Pr√©sence des arbitres")
            commentaires = {}
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Pr√©nom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")
                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Pr√©sent", "Absent excus√©", "Absent non excus√©"], key=f"s_statut_{i}")
                commentaire = col2.text_input("Commentaire individuel", key=f"s_comment_{i}")
                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()


            observations = st.text_area("Observations g√©n√©rales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_stage:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_stage]

                    rass.append({
                        "Nom": nom_stage,
                        "Type": "Stage",
                        "Date d√©but": date_debut.strftime("%d/%m/%Y"),
                        "Date fin": date_fin.strftime("%d/%m/%Y"),
                        "Statut": statuts[f"{a['Pr√©nom']} {a['Nom']}"],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Pr√©nom']} {a['Nom']}", "")

                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("Stage enregistr√© avec succ√®s.")

    elif type_rass == "Test physique":
        commentaires = {}
        with st.form("form_test_physique"):
            nom_test = st.text_input("Nom du test physique")
            date_test = st.date_input("Date du test")
            statuts = {}
            st.markdown("### R√©sultat des arbitres")
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Pr√©nom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")

                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Pr√©sent", "Absent excus√©", "Absent non excus√©"], key=f"r_statut_{i}")
                commentaire = col2.text_input("Commentaire", key=f"r_comment_{i}")

                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()


            observations = st.text_area("Observations g√©n√©rales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_test:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_test]
                    nom_complet = f"{a['Pr√©nom']} {a['Nom']}"
                    rass.append({
                        "Nom": nom_test,
                        "Type": "Test physique",
                        "Date": date_test.strftime("%d/%m/%Y"),
                        "Statut": statuts[nom_complet],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Pr√©nom']} {a['Nom']}", "")
                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("Test physique enregistr√© avec succ√®s.")


elif action == "üìä R√©capitulatif des rassemblements":
    st.subheader("üìä R√©capitulatif des rassemblements")

    # Chargement des rassemblements
    rassemblements = {}
    for i, arbitre in enumerate(st.session_state["far_arbitres"]):
        try:
            rass_list = json.loads(arbitre.get("Rassemblements", "")) if arbitre.get("Rassemblements") else []
        except:
            rass_list = []

        for r in rass_list:
            nom_rass = r.get("Nom")
            if nom_rass not in rassemblements:
                rassemblements[nom_rass] = {
                    "Type": r.get("Type"),
                    "Date_sort": r.get("Date d√©but", r.get("Date", "")),  # pour tri
                    "Dates": f"{r.get('Date d√©but', r.get('Date', ''))} ‚Üí {r.get('Date fin', '')}".strip(" ‚Üí"),
                    "Pr√©sences": []
                }
            rassemblements[nom_rass]["Pr√©sences"].append((arbitre["Pr√©nom"], arbitre["Nom"], r.get("Statut", "Inconnu")))

    if not rassemblements:
        st.info("Aucun rassemblement enregistr√©.")
    else:
        # Filtrage par type
        all_types = ["Tous"] + sorted(list(set([r["Type"] for r in rassemblements.values()])))
        selected_type = st.selectbox("Filtrer par type", all_types)

        # Tri des rassemblements
        sorted_rass = sorted(
            rassemblements.items(),
            key=lambda x: datetime.strptime(x[1]["Date_sort"], "%d/%m/%Y") if x[1]["Date_sort"] else datetime.min
        )

        for nom, infos in sorted_rass:
            if selected_type != "Tous" and infos["Type"] != selected_type:
                continue

            statuts_valides = ["pr√©sent", "r√©ussi", "√©chec", "√©chec partiel"]
            present = sum(1 for p in infos["Pr√©sences"] if p[2].lower().strip() in statuts_valides)
            absents_exc = sum(1 for p in infos["Pr√©sences"] if p[2].lower().strip() == "absent excus√©")
            absents_non = sum(1 for p in infos["Pr√©sences"] if p[2].lower().strip() == "absent non excus√©")
            total = len(infos["Pr√©sences"])


            with st.expander(f"üìå {nom} ({infos['Type']}) ‚Äî {infos['Dates']}"):
                st.markdown(f"üë• {total} arbitres concern√©s")
                st.markdown(f"‚úÖ {present} pr√©sent(s)")
                st.markdown(f"‚ùå {absents_exc} absent(s) excus√©(s), {absents_non} absent(s) non excus√©(s)")

                st.markdown("**Participants :**")
                for prenom, nom_arbitre, statut in infos["Pr√©sences"]:
                    st.markdown(f"- {prenom} {nom_arbitre} ‚Üí *{statut}*")

                if st.button(f"üóëÔ∏è Supprimer ce rassemblement", key=f"delete_rass_{nom}"):
                    for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                        rass = json.loads(arbitre.get("Rassemblements", "")) if arbitre.get("Rassemblements") else []
                        rass = [r for r in rass if r.get("Nom") != nom]
                        st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                    save_arbitres(st.session_state["far_arbitres"])
                    st.success(f"Rassemblement '{nom}' supprim√©.")
                    st.rerun()

elif action == "üõë Ajouter des manquements":
    st.subheader("üõë Ajouter un manquement √† un arbitre")

    # Liste des rassemblements disponibles
    all_rass_names = set()
    for a in st.session_state["far_arbitres"]:
        rass_list = json.loads(a.get("Rassemblements", "") or "[]")
        for r in rass_list:
            all_rass_names.add(r["Nom"])
    rass_names = sorted(list(all_rass_names))

    with st.form("form_ajout_manquement"):
        # S√©lection de l‚Äôarbitre
        arbitres_dict = {f"{a['Pr√©nom']} {a['Nom']}": i for i, a in enumerate(st.session_state["far_arbitres"])}
        nom_sel = st.selectbox("üë§ S√©lectionner un arbitre", [""] + list(arbitres_dict.keys()))

        # Type de manquement
        type_manq = st.selectbox("üìå Type de manquement", [
            "",  # Option vide
            "Non-r√©ponse √† un Google Form",
            "Absence non excus√©e",
            "Livrable non rendu",
            "Livrable rendu en retard",
            "Retard",
            "Indisponibilit√© tardive",
            "Autre"
        ])

        # Date (toujours demand√©)
        date = st.date_input("üìÖ Date du manquement")

        # Initialisation des variables
        detail = ""
        commentaire = ""

        # Champ commentaire (optionnel)
        commentaire = st.text_area("üóíÔ∏è Commentaire (optionnel)")

        submit = st.form_submit_button("‚úÖ Ajouter le manquement")

        if submit:
            if not nom_sel:
                st.warning("Veuillez s√©lectionner un arbitre.")
            elif not type_manq:
                st.warning("Veuillez s√©lectionner un type de manquement.")
            else:
                i = arbitres_dict[nom_sel]
                try:
                    entry = {
                        "Type": type_manq,
                        "Date": date.strftime("%d/%m/%Y"),
                        "D√©tail": detail
                    }
                    if commentaire.strip():
                        entry["Commentaire"] = commentaire.strip()

                    raw = st.session_state["far_arbitres"][i].get("Manquements", "")
                    if not isinstance(raw, str):
                        raw = "" if raw is None or (isinstance(raw, float) and math.isnan(raw)) else str(raw)
                    mqs = json.loads(raw or "[]")
                    mqs.append(entry)
                    st.session_state["far_arbitres"][i]["Manquements"] = json.dumps(mqs)
                    save_arbitres(st.session_state["far_arbitres"])
                    st.success(f"Manquement ajout√© pour {nom_sel}")
                except Exception as e:
                    st.error(f"Erreur lors de l'enregistrement : {e}")



elif action == "üìâ R√©capitulatif des manquements":
    st.subheader("üìâ R√©capitulatif des manquements")

    tous_manquements = []

    for i, a in enumerate(st.session_state["far_arbitres"]):
        try:
            mqs = json.loads(a.get("Manquements", "")) if a.get("Manquements") else []
        except:
            mqs = []

        for m in mqs:
            ligne = {
                "Index": i,
                "Nom complet": f"{a['Pr√©nom']} {a['Nom']}",
                "Type": m.get("Type", "Inconnu"),
                "D√©tail": m.get("D√©tail", ""),
                "Date": m.get("Date", "")
            }
            tous_manquements.append(ligne)

    if not tous_manquements:
        st.info("Aucun manquement enregistr√©.")
    else:
        df_mqs = pd.DataFrame(tous_manquements)

        types_dispo = ["Tous"] + sorted(df_mqs["Type"].unique())
        filtre_type = st.selectbox("Filtrer par type de manquement", types_dispo)

        if filtre_type != "Tous":
            df_mqs = df_mqs[df_mqs["Type"] == filtre_type]

        df_mqs = df_mqs.sort_values(by="Date", ascending=False)

        for idx, row in df_mqs.iterrows():
            with st.expander(f"üö´ {row['Nom complet']} ‚Äì {row['Type']}"):
                st.markdown(f"üìÖ **Date :** {row['Date']}")
                if row["D√©tail"]:
                    st.markdown(f"üìù **D√©tail :** {row['D√©tail']}")

                if st.button("üóëÔ∏è Supprimer ce manquement", key=f"del_manquement_{idx}"):
                    i = row["Index"]
                    try:
                        mqs = json.loads(st.session_state["far_arbitres"][i].get("Manquements", "")) or []
                        mqs = [m for m in mqs if not (
                            m.get("Type") == row["Type"] and
                            m.get("Date") == row["Date"] and
                            m.get("D√©tail", "") == row["D√©tail"]
                        )]
                        st.session_state["far_arbitres"][i]["Manquements"] = json.dumps(mqs)
                        save_arbitres(st.session_state["far_arbitres"])
                        st.success("Manquement supprim√©.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erreur : {e}")

elif action == "üìù Saisie des examens":
    st.subheader("üìù Saisie des r√©sultats d'examen")

    with st.form("form_saisie_examen"):
        nom_examen = st.text_input("Nom de l'examen")
        date_examen = st.date_input("Date", value=date.today())
        is_probatoire = st.checkbox("‚úÖ Examen probatoire")

        notes = {}
        statuts = {}

        if not is_probatoire:
            note_max = st.number_input("Note maximale", min_value=1.0, max_value=100.0, value=20.0, step=0.5)

        for i, arbitre in enumerate(st.session_state["far_arbitres"]):
            nom_complet = f"{arbitre['Pr√©nom']} {arbitre['Nom']}"
            st.markdown(f"#### {nom_complet}")

            statut = st.selectbox("Statut", ["Pr√©sent", "Absent excus√©", "Absent non excus√©"], key=f"statut_examen_{i}")
            statuts[i] = statut

            if statut != "Pr√©sent":
                notes[i] = {"Absent": True, "Statut": statut}
                st.warning(f"Statut : {statut}")
            else:
                if is_probatoire:
                    qcm = st.number_input("Questionnaire th√©orique (sur 50)", min_value=0.0, max_value=50.0, step=0.5, key=f"qcm_examen_{i}")
                    video = st.number_input("Test vid√©o (sur 30)", min_value=0.0, max_value=30.0, step=0.5, key=f"video_examen_{i}")
                    rapport = st.number_input("Rapport disciplinaire (sur 20)", min_value=0.0, max_value=20.0, step=0.5, key=f"rapport_examen_{i}")
                    total = round(qcm + video + rapport, 2)
                    st.markdown(f"üìä **Total : {total}/100**")
                    notes[i] = {
                        "QCM": qcm,
                        "Vid√©o": video,
                        "Rapport": rapport,
                        "Total": total,
                        "Absent": False,
                        "Statut": statut
                    }
                else:
                    note_unique = st.number_input(f"Note (sur {note_max})", min_value=0.0, max_value=note_max, step=0.5, key=f"note_unique_{i}")
                    notes[i] = {
                        "Note": note_unique,
                        "Sur": note_max,
                        "Absent": False,
                        "Statut": statut
                    }

        submit_notes = st.form_submit_button("Enregistrer les notes")

        if submit_notes and nom_examen:

            # Supprimer l'examen existant (nom_examen) pour tous les arbitres avant d'enregistrer
            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                raw = arbitre.get("Examens", "")
                try:
                    examens = json.loads(raw) if isinstance(raw, str) and raw.strip() else []
                except:
                    examens = []

                new_examens = [e for e in examens if e.get("Nom") != nom_examen]
                st.session_state["far_arbitres"][i]["Examens"] = json.dumps(new_examens)

            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                raw = arbitre.get("Examens", "")
                try:
                    examens = json.loads(raw) if isinstance(raw, str) and raw.strip() else []
                except:
                    examens = []

                # Supprimer un √©ventuel doublon
                examens = [e for e in examens if e.get("Nom") != nom_examen]

                data = notes.get(i, {})
                examen_dict = {
                    "Nom": nom_examen,
                    "Type": "Examen",
                    "Date": date_examen.strftime("%d/%m/%Y"),
                    "Statut": data.get("Statut", "Pr√©sent"),
                    "Type examen": "Probatoire" if is_probatoire else "Classique"
                }

                if not data.get("Absent", False):
                    if is_probatoire:
                        examen_dict.update({
                            "QCM": data.get("QCM"),
                            "Video": data.get("Vid√©o"),
                            "Rapport": data.get("Rapport"),
                            "Note": data.get("Total"),
                            "Sur": 100
                        })
                    else:
                        examen_dict.update({
                            "Note": data.get("Note"),
                            "Sur": data.get("Sur")
                        })

                examens.append(examen_dict)
                st.session_state["far_arbitres"][i]["Examens"] = json.dumps(examens)

            save_arbitres(st.session_state["far_arbitres"])
            st.success("‚úÖ R√©sultats d'examen enregistr√©s avec succ√®s.")


elif action == "üìä R√©capitulatif des examens":
    st.subheader("üìä R√©capitulatif des examens")

    import matplotlib.pyplot as plt
    from collections import defaultdict

    examens = []
    for i, a in enumerate(st.session_state["far_arbitres"]):
        try:
            exam_list = json.loads(a.get("Examens", "") or "[]")
        except:
            exam_list = []

        for r in exam_list:
            examens.append({
                "Index": i,
                "Nom complet": f"{a['Pr√©nom']} {a['Nom']}",
                "Cat√©gorie": a.get("Cat√©gorie", ""),
                "Type": r.get("Type examen", "Classique"),
                "Date": r.get("Date"),
                "Nom": r.get("Nom"),
                "Statut": r.get("Statut"),
                "Note": float(r.get("Note", 0)),
                "Sur": float(r.get("Sur", 100)),
                "QCM": r.get("QCM", None),
                "Video": r.get("Video", None),
                "Rapport": r.get("Rapport", None),
            })




    if not examens:
        st.info("Aucun examen enregistr√©.")
        st.stop()

    df_examens = pd.DataFrame(examens)

    # Statistiques globales
    st.markdown("### üìà Statistiques globales")

    nb_exams_uniques = df_examens["Nom"].nunique()
    nb_absents = df_examens[df_examens["Statut"] != "Pr√©sent"].groupby("Nom").ngroups
    taux_presence = 100 * (nb_exams_uniques - nb_absents) / nb_exams_uniques if nb_exams_uniques else 0

    st.markdown(f"- Nombre total d'examens : **{nb_exams_uniques}**")
    st.markdown(f"- Nombre d'examens avec au moins un absent : **{nb_absents}**")
    st.markdown(f"- Taux de pr√©sence globale : **{taux_presence:.1f}%**")


    # Graphiques
    st.markdown("### üìä Graphiques")

    # Moyenne par arbitre
    df_notes = df_examens[df_examens["Statut"] == "Pr√©sent"].copy()
    df_notes["Pourcentage"] = df_notes["Note"] / df_notes["Sur"] * 100

    df_examens["Pourcentage"] = 100 * df_examens["Note"] / df_examens["Sur"]
    moyennes = df_notes.groupby("Nom complet")["Pourcentage"].mean().sort_values()
    totaux = df_examens.groupby("Nom complet")["Pourcentage"].sum().sort_values()

    fig1, ax1 = plt.subplots(figsize=(8, 4))
    moyennes.plot(kind="barh", ax=ax1)
    ax1.set_title("Moyenne % par arbitre (pr√©sents uniquement)")
    st.pyplot(fig1)

    fig2, ax2 = plt.subplots(figsize=(8, 4))
    totaux.plot(kind="barh", ax=ax2)
    ax2.set_title("Total % par arbitre (absences p√©nalisantes)")
    st.pyplot(fig2)

    # Classement par cat√©gorie
    st.markdown("### üèÜ Classement par cat√©gorie")

    cat_options = df_examens["Cat√©gorie"].dropna().unique().tolist()
    selected_cat = st.selectbox("Filtrer par cat√©gorie", ["Toutes"] + sorted(cat_options))

    if selected_cat != "Toutes":
        df_notes = df_notes[df_notes["Cat√©gorie"] == selected_cat]
        df_examens = df_examens[df_examens["Cat√©gorie"] == selected_cat]

    classement_moyenne = df_notes.groupby("Nom complet")["Pourcentage"].mean().sort_values(ascending=False)
    classement_total = df_examens.groupby("Nom complet")["Pourcentage"].sum().sort_values(ascending=False)

    st.markdown("#### Classement par moyenne")
    for i, (nom, score) in enumerate(classement_moyenne.items(), 1):
        st.markdown(f"{i}. **{nom}** ‚Äì {score:.1f}%")

    st.markdown("#### Classement par total")
    for i, (nom, score) in enumerate(classement_total.items(), 1):
        st.markdown(f"{i}. **{nom}** ‚Äì {score:.1f} points cumul√©s")

    # Table compl√®te
    st.markdown("### üìã D√©tails des examens")
    # üéØ Filtrer par nom d'examen
    examens_disponibles = df_examens["Nom"].unique().tolist()
    selected_exam = st.selectbox("üìù Filtrer par examen", ["Tous"] + sorted(examens_disponibles))

    df_display = df_examens.copy()
    df_display["Note (%)"] = df_display["Pourcentage"].map(lambda x: f"{x:.1f}%")

    # üß† S√©lection des colonnes selon type d‚Äôexamen
    colonnes_communes = ["Nom", "Date", "Nom complet", "Cat√©gorie", "Type", "Statut"]

    if selected_exam != "Tous":
        df_display = df_display[df_display["Nom"] == selected_exam]

    # Teste si examen probatoire
    is_probatoire = (
        df_display["Type"].iloc[0] == "Probatoire"
        if not df_display.empty
        else False
    )

    if is_probatoire:
        colonnes = colonnes_communes + [c for c in ["QCM", "Video", "Rapport", "Note", "Sur", "Note (%)"] if c in df_display.columns]
    else:
        colonnes = colonnes_communes + [c for c in ["Note", "Sur", "Note (%)"] if c in df_display.columns]

    # Filtrage des colonnes disponibles
    df_display = df_display[[col for col in colonnes if col in df_display.columns]]

    st.dataframe(df_display, use_container_width=True)

    if selected_exam != "Tous":
        df_exam_unique = df_examens[df_examens["Nom"] == selected_exam].copy()

        st.markdown(f"## üìÑ R√©sultats d√©taill√©s pour l'examen : {selected_exam}")

        # üî¢ Classement
        classement = df_exam_unique[df_exam_unique["Statut"] == "Pr√©sent"].copy()
        classement["Pourcentage"] = 100 * classement["Note"] / classement["Sur"]
        classement = classement.sort_values(by="Pourcentage", ascending=False)

        st.markdown("### üèÖ Classement des arbitres (pr√©sents uniquement)")
        for rang, (_, row) in enumerate(classement.iterrows(), start=1):
            st.markdown(f"{rang}. **{row['Nom complet']}** ‚Äì {row['Note']:.1f} / {row['Sur']}")


        # üìä Graphique
        import matplotlib.pyplot as plt

        st.markdown("### üìä Graphique des r√©sultats")
        fig, ax = plt.subplots(figsize=(10, 4))
        classement.set_index("Nom complet")["Pourcentage"].plot(kind="barh", ax=ax)
        ax.invert_yaxis()
        ax.set_xlabel("Note en %")
        ax.set_title(f"R√©sultats √† l'examen ¬´ {selected_exam} ¬ª")
        st.pyplot(fig)

        # üßæ Tableau d√©taill√© exportable
        st.markdown("### üìã Tableau des r√©sultats")
        colonnes_affichage = ["Nom complet", "Cat√©gorie", "Statut", "Type", "Note", "Sur", "Pourcentage"]

        if df_exam_unique["Type"].iloc[0] == "Probatoire":
            colonnes_affichage = ["Nom complet", "Cat√©gorie", "Statut", "QCM", "Video", "Rapport", "Note", "Sur", "Pourcentage"]

        df_export = df_exam_unique.copy()
        df_export["Pourcentage"] = df_export["Pourcentage"].map(lambda x: round(x, 1))
        df_export = df_export.reset_index(drop=True)  # pour √©viter d√©calages d'index
        df_export.insert(0, "Rang", df_export.index + 1)


        st.dataframe(df_export[colonnes_affichage], use_container_width=True)

        # üì• Bouton de t√©l√©chargement
        from io import BytesIO

        buffer = BytesIO()
        df_export[colonnes_affichage].to_excel(buffer, index=False)
        buffer.seek(0)

        st.download_button(
            label="üì• T√©l√©charger les r√©sultats",
            data=buffer,
            file_name=f"R√©sultats_{selected_exam}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    st.markdown("### üßπ Supprimer un examen")

    noms_examens_dispo = df_examens["Nom"].unique().tolist()
    examen_a_supprimer = st.selectbox("üìå S√©lectionnez un examen √† supprimer", [""] + noms_examens_dispo)

    if examen_a_supprimer:
        if st.button("üóëÔ∏è Supprimer cet examen pour tous les arbitres"):
            nb_suppr = 0
            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                examens = json.loads(arbitre.get("Examens", "") or "[]")
                new_examens = [e for e in examens if e.get("Nom") != examen_a_supprimer]
                if len(new_examens) != len(examens):  # s'il y a eu une suppression
                    st.session_state["far_arbitres"][i]["Examens"] = json.dumps(new_examens)
                    nb_suppr += 1

            if nb_suppr > 0:
                save_arbitres(st.session_state["far_arbitres"])
                st.success(f"‚úÖ Examen '{examen_a_supprimer}' supprim√© pour {nb_suppr} arbitre(s).")
                st.rerun()
            else:
                st.info("Aucun arbitre n'√©tait concern√© par cet examen.")

if action == "üìé D√©poser un rapport d'observation":
    st.subheader("üìé D√©poser un rapport pour un arbitre")

    # S√©lection de l‚Äôarbitre
    arbitres_dict = {f"{a['Pr√©nom']} {a['Nom']}": i for i, a in enumerate(st.session_state["far_arbitres"])}

    with st.form("upload_rapport_form"):
        nom_sel = st.selectbox("üë§ S√©lectionner un arbitre", [""] + list(arbitres_dict.keys()))
        uploaded_file = st.file_uploader("üìÑ D√©poser un rapport (PDF, Word...)", type=["pdf", "docx", "doc"])
        url_partage = upload_to_drive("temp_upload.pdf", uploaded_file.name, parent_folder_id=folder_id)

        submit = st.form_submit_button("‚úÖ Enregistrer le rapport")

    if submit and nom_sel and uploaded_file:
        # Sauvegarder temporairement le fichier
        filename_temp = f"temp_{uuid.uuid4().hex}.pdf"
        with open(filename_temp, "wb") as f:
            f.write(uploaded_file.getbuffer())


        os.remove(filename_temp)

        # Uploader sur Drive
        url_partage = upload_to_drive("temp_upload.pdf", uploaded_file.name, parent_folder_id=folder_id)

        # Sauvegarder dans l‚Äôarbitre
        rapports = json.loads(a.get("Rapports", "[]"))
        rapports.append({
            "nom_original": uploaded_file.name,
            "url": url_partage
        })
        a["Rapports"] = json.dumps(rapports)
        save_arbitres(st.session_state["far_arbitres"])
        st.success("‚úÖ Rapport envoy√© sur Google Drive.")
        st.rerun()




elif action == "üë§ Fiche arbitre":
    st.subheader("üë§ Rechercher un arbitre")
    noms = [f"{a['Pr√©nom']} {a['Nom']}" for a in st.session_state["far_arbitres"]]
    sel = st.selectbox("S√©lectionnez un arbitre", [""] + sorted(noms))
    
    if sel:
        a = next(a for a in st.session_state["far_arbitres"] if f"{a['Pr√©nom']} {a['Nom']}" == sel)

        st.markdown(f"### {sel} ‚Äî {a.get('Cat√©gorie')} ‚Äî √Çge¬†: {a.get('√Çge')}")
        st.markdown(f"- üìÖ Date de naissance¬†: {a.get('Date de naissance')}")
        st.markdown(f"- üè† Club¬†: {a.get('Club')}")
        tel = str(a.get("T√©l√©phone", "")).strip()
        if tel and not tel.startswith("0") and len(tel) == 9:
            tel = "0" + tel
        email = a.get("Email", "‚Äî")

        st.markdown(f"- üìû Tel¬†: {tel}  |  ‚úâÔ∏è Email¬†: {email}")

        # === Rassemblements ===
        rass = json.loads(a.get("Rassemblements", "") or "[]")
        if rass:
            st.markdown("### üìã Rassemblements")

            # Convertir en DataFrame
            df_rass = pd.DataFrame(rass)

            # Fusionner Date d√©but et Date fin en une date d‚Äôaffichage
            df_rass["Date"] = df_rass.apply(
                lambda row: row["Date"] if "Date" in row and pd.notna(row["Date"])
                else row.get("Date d√©but", ""), axis=1
            )

           

            # Convertir en datetime pour tri
            df_rass["Date_dt"] = pd.to_datetime(df_rass["Date"], format="%d/%m/%Y", errors="coerce")


            # Regrouper par type
            for type_rass in sorted(df_rass["Type"].dropna().unique()):
                st.markdown(f"#### üóÇÔ∏è {type_rass}")

                df_sous = df_rass[df_rass["Type"] == type_rass].copy()
                df_sous = df_sous.sort_values("Date_dt")

                # Affichage propre
                df_sous = df_sous.rename(columns={"Observations individuelles": "Observation"})
                for col in ["Nom", "Date", "Type", "Statut", "Observation"]:
                    if col not in df_sous.columns:
                        df_sous[col] = ""

                df_sous_affiche = df_sous[["Nom", "Date", "Type", "Statut", "Observation"]].reset_index(drop=True)
                df_sous_affiche.index = df_sous_affiche.index + 1
                st.dataframe(df_sous_affiche, use_container_width=True)


        # === Examens ===
            exam = json.loads(a.get("Examens", "") or "[]")
            if exam:
                st.markdown("#### üß† Examens")
                for e in exam:
                    if "Note" not in e and "Total" in e:
                        e["Note"] = e["Total"]
                    if "Type examen" not in e:
                        e["Type examen"] = ""
                    if e.get("Type examen") != "Probatoire":
                        e["QCM"] = e["Video"] = e["Rapport"] = ""

                df_exam = pd.DataFrame(exam)
                colonnes = ["Nom", "Date", "Type examen", "Statut", "Note", "Sur", "QCM", "Video", "Rapport"]
                df_exam = df_exam[[c for c in colonnes if c in df_exam.columns]]
                df_exam.index = df_exam.index + 1
                st.dataframe(df_exam, use_container_width=True)


                # === Position globale dans la cat√©gorie ===
                nom_complet = f"{a['Pr√©nom']} {a['Nom']}"
                cat = a.get("Cat√©gorie")
                total_points = 0

                # Total de l'arbitre courant
                exam = json.loads(a.get("Examens", "") or "[]")
                for e in exam:
                    try:
                        total_points += float(e.get("Note", e.get("Total", 0)))
                    except:
                        continue

                # Calcul du classement global
                classement = []
                for autre in st.session_state["far_arbitres"]:
                    if autre.get("Cat√©gorie") != cat:
                        continue
                    try:
                        raw = autre.get("Examens", "")
                        autres_exams = json.loads(raw if isinstance(raw, str) else "")
                        total = sum(float(e.get("Note", e.get("Total", 0))) for e in autres_exams)
                        nom_autre = f"{autre.get('Pr√©nom')} {autre.get('Nom')}"
                        classement.append((nom_autre, total))
                    except:
                        continue

                classement.sort(key=lambda x: x[1], reverse=True)
                rang = next((i + 1 for i, (n, _) in enumerate(classement) if n == nom_complet), None)
                total_cat = len(classement)

                if rang:
                    st.markdown(f"### üìä Position globale dans la cat√©gorie **{cat}** : {rang}·µâ sur {total_cat}")



        # === Manquements ===
        val = a.get("Manquements", "")
        try:
            mqs = json.loads(val if isinstance(val, str) else "")
        except:
            mqs = []

        if mqs:
            st.markdown("#### üö´ Manquements")
            df_manq = pd.DataFrame(mqs)
            colonnes = ["Date", "Type", "D√©tail", "Commentaire"]
            df_manq = df_manq[[c for c in colonnes if c in df_manq.columns]]
            df_manq.index = df_manq.index + 1
            st.dataframe(df_manq, use_container_width=True)
        

        # === Rapports d'observation ===
        st.markdown("### üìÅ Rapports d'observation")
        raw_rapports = json.loads(a.get("Rapports", "[]"))

        # Compatibilit√© : transformer ancienne liste de strings en dicts
        rapports = []
        for r in raw_rapports:
            if isinstance(r, str):
                rapports.append({"fichier": r, "nom_original": r})
            else:
                rapports.append(r)

        if rapports:
            for i, rapport in enumerate(rapports):
                nom_affiche = rapport.get("nom_original", f"rapport_{i+1}")

                # === Cas 1 : lien Google Drive
                if "url" in rapport:
                    url = rapport["url"]
                    col1, col2 = st.columns([6, 1])
                    with col1:
                        st.markdown(f"üìÑ [{nom_affiche}]({url})", unsafe_allow_html=True)
                    with col2:
                        if st.button("üóëÔ∏è", key=f"del_drive_{i}"):
                            rapports.pop(i)
                            a["Rapports"] = json.dumps(rapports)
                            save_arbitres(st.session_state["far_arbitres"])
                            st.success(f"Rapport supprim√© : {nom_affiche}")
                            st.rerun()

                # === Cas 2 : fichier local
                elif "fichier" in rapport:
                    nom_fichier = rapport["fichier"]
                    path = os.path.join("rapports", nom_fichier)
                    if os.path.exists(path):
                        col1, col2 = st.columns([6, 1])
                        with col1:
                            with open(path, "rb") as f:
                                st.download_button(
                                    label=f"üì• T√©l√©charger {nom_affiche}",
                                    data=f,
                                    file_name=nom_fichier,
                                    mime="application/octet-stream",
                                    key=f"dl_{nom_fichier}_{i}"
                                )
                        with col2:
                            if st.button("üóëÔ∏è", key=f"del_local_{i}"):
                                try:
                                    os.remove(path)
                                except Exception as e:
                                    st.warning(f"Impossible de supprimer le fichier : {e}")
                                rapports.pop(i)
                                a["Rapports"] = json.dumps(rapports)
                                save_arbitres(st.session_state["far_arbitres"])
                                st.success(f"Rapport supprim√© : {nom_affiche}")
                                st.rerun()
        else:
            st.info("Aucun rapport enregistr√©.")





        # === Boutons Word ===
        from docx import Document
        from io import BytesIO
        from docx.shared import Pt
        from datetime import datetime
        import json

        def create_doc_for_arbitre(a, arbitres_liste):
            doc = Document()
            nom_complet = f"{a['Pr√©nom']} {a['Nom']}"

            # === En-t√™te ===
            doc.add_heading(nom_complet, level=1)
            doc.add_paragraph(f"Cat√©gorie : {a.get('Cat√©gorie', '‚Äî')}")
            doc.add_paragraph(f"√Çge : {a.get('√Çge', '‚Äî')}")
            doc.add_paragraph(f"N√©(e) le : {a.get('Date de naissance', '‚Äî')}")
            doc.add_paragraph(f"Club : {a.get('Club', '‚Äî')}")
            doc.add_paragraph(f"T√©l√©phone : {str(a.get('T√©l√©phone', '‚Äî'))}")
            doc.add_paragraph(f"Email : {a.get('Email', '‚Äî')}")

            # === Examens ===
            exams_raw = a.get("Examens", "")
            try:
                exams = json.loads(exams_raw if isinstance(exams_raw, str) else "")
            except:
                exams = []

            if exams:
                doc.add_heading("Examens", level=2)
                table = doc.add_table(rows=1, cols=8)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Nom"
                hdr[1].text = "Date"
                hdr[2].text = "Type"
                hdr[3].text = "Statut"
                hdr[4].text = "Note"
                hdr[5].text = "QCM"
                hdr[6].text = "Vid√©o"
                hdr[7].text = "Rapport"

                total_points = 0
                for e in exams:
                    note = float(e.get("Note", e.get("Total", 0)))
                    total_points += note

                    row = table.add_row().cells
                    row[0].text = e.get("Nom", "?")
                    row[1].text = e.get("Date", "?")
                    row[2].text = e.get("Type examen", "")
                    row[3].text = e.get("Statut", "")
                    row[4].text = f"{note} / {e.get('Sur', 100)}"
                    row[5].text = str(e.get("QCM", "")) if e.get("Type examen") == "Probatoire" else ""
                    row[6].text = str(e.get("Video", "")) if e.get("Type examen") == "Probatoire" else ""
                    row[7].text = str(e.get("Rapport", "")) if e.get("Type examen") == "Probatoire" else ""

            else:
                total_points = 0

            # === Rassemblements ===
            rass_raw = a.get("Rassemblements", "")
            try:
                rass = json.loads(rass_raw if isinstance(rass_raw, str) else "")
            except:
                rass = []

            if rass:
                rass.sort(key=lambda r: datetime.strptime(r.get("Date", "01/01/1900"), "%d/%m/%Y"))
                doc.add_heading("Rassemblements", level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Nom"
                hdr[1].text = "Date"
                hdr[2].text = "Type"
                hdr[3].text = "Statut"
                hdr[4].text = "Observation individuelle"

                for r in rass:
                    row = table.add_row().cells
                    row[0].text = r.get("Nom", "?")
                    row[1].text = r.get("Date", "?")
                    row[2].text = r.get("Type", "")
                    row[3].text = r.get("Statut", "")
                    row[4].text = r.get("Observations individuelles", "")

            # === Manquements ===
            val = a.get("Manquements", "")
            try:
                mqs = json.loads(val if isinstance(val, str) else "")
            except:
                mqs = []

            if mqs:
                doc.add_heading("Manquements", level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Date"
                hdr[1].text = "Type"
                hdr[2].text = "D√©tail"

                for m in mqs:
                    row = table.add_row().cells
                    row[0].text = m.get("Date", "")
                    row[1].text = m.get("Type", "")
                    row[2].text = m.get("D√©tail", "")

            # === Position globale dans la cat√©gorie ===
            cat = a.get("Cat√©gorie")
            if cat and total_points > 0:
                classement = []
                for autre in arbitres_liste:
                    if autre.get("Cat√©gorie") != cat:
                        continue
                    try:
                        raw = autre.get("Examens", "")
                        autres_exams = json.loads(raw if isinstance(raw, str) else "")
                        total = sum(float(e.get("Note", e.get("Total", 0))) for e in autres_exams)
                        nom_autre = f"{autre.get('Pr√©nom')} {autre.get('Nom')}"
                        classement.append((nom_autre, total))
                    except:
                        continue

                classement.sort(key=lambda x: x[1], reverse=True)
                rang = next((i+1 for i, (n, _) in enumerate(classement) if n == nom_complet), None)
                total_cat = len(classement)

                if rang:
                    doc.add_paragraph(f"üìä Position globale dans la cat√©gorie **{cat}** : {rang}·µâ sur {total_cat}")

            return doc


        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            if st.button("üìÑ G√©n√©rer fiche Word"):
                doc = create_doc_for_arbitre(a, st.session_state["far_arbitres"])
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "T√©l√©charger fiche Word",
                    data=buffer,
                    file_name=f"{sel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        with col2:
            if st.button("üìÅ G√©n√©rer fiches Word (tous les arbitres)"):
                doc_all = Document()
                for arbitre in st.session_state["far_arbitres"]:
                    fiche = create_doc_for_arbitre(arbitre, st.session_state["far_arbitres"])
                    for p in fiche.paragraphs:
                        doc_all.add_paragraph(p.text)
                    doc_all.add_page_break()

                buffer = BytesIO()
                doc_all.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "T√©l√©charger fiches (tous)",
                    data=buffer,
                    file_name="fiches_arbitres.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )






# === AFFICHAGE LISTE ===
st.markdown("---")
st.subheader("üìã Liste des arbitres")
for a in st.session_state["far_arbitres"]:
    st.markdown(f"- **{a['Pr√©nom']} {a['Nom']}** | {a['Cat√©gorie']} | √Çge : {a.get('√Çge', 'N/A')}")
